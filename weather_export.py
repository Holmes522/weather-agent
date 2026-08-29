"""天气快照导出为常见办公文档格式。"""

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PDF_FONT_NAME = "STSong-Light"
pdfmetrics.registerFont(UnicodeCIDFont(PDF_FONT_NAME))

EXPORT_ACTION_RE = re.compile(r"导出|保存|生成|输出|下载|存储|做成|整理成")
FORMAT_PATTERNS = (
    ("docx", re.compile(r"docx?|word|文档", re.IGNORECASE)),
    ("xlsx", re.compile(r"xlsx?|excel|execl|表格", re.IGNORECASE)),
    ("pdf", re.compile(r"pdf", re.IGNORECASE)),
    ("md", re.compile(r"markdown|md", re.IGNORECASE)),
)
MAX_EXPORT_RECORDS = 35


@dataclass(frozen=True)
class ExportRequest:
    requested: bool
    format: Optional[str] = None


@dataclass(frozen=True)
class WeatherSnapshot:
    city: str
    date_label: str
    provider: str
    temperature_c: float
    condition: str
    humidity_percent: int
    wind_speed_mps: float
    rain_expected: bool
    advice: Optional[str] = None


@dataclass(frozen=True)
class WeatherReportContext:
    """报告级行程信息；不包含密钥或用户自由路径。"""

    origin: Optional[str] = None
    total_days: int = 0
    ordered: bool = False


@dataclass(frozen=True)
class WeatherArtifact:
    filename: str
    mimetype: str
    content: bytes


def parse_export_request(message: str) -> ExportRequest:
    """识别显式导出动作；只出现文件扩展名不视为导出请求。"""

    if not isinstance(message, str) or not EXPORT_ACTION_RE.search(message):
        return ExportRequest(False)
    for format_name, pattern in FORMAT_PATTERNS:
        if pattern.search(message):
            return ExportRequest(True, format_name)
    return ExportRequest(True)


def weather_query_text(message: str) -> str:
    """只移除导出动作和格式，保留其后的行程、出发地和停留信息。"""

    cleaned = re.sub(
        r"(?:结果)?\s*(?:导出|保存|生成|输出|下载|存储|做成|整理成)"
        r"\s*(?:来|为|成)?\s*"
        r"(?:docx?|word|文档|xlsx?|excel|execl|表格|pdf|markdown|md|文件|表)?",
        " ",
        message,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" ，,。；;")
    return re.sub(r"^(?:请)?把", "", cleaned).strip()


def build_weather_document(
    snapshots: Sequence[WeatherSnapshot],
    format_name: str,
    report_context: Optional[WeatherReportContext] = None,
) -> WeatherArtifact:
    """生成有界、无宏无公式的天气报告字节。"""

    records = tuple(snapshots)[:MAX_EXPORT_RECORDS]
    if not records:
        raise ValueError("weather snapshots are required")
    builders = {
        "md": _build_markdown,
        "docx": _build_docx,
        "xlsx": _build_xlsx,
        "pdf": _build_pdf,
    }
    builder = builders.get(format_name)
    if builder is None:
        raise ValueError("unsupported export format")
    return builder(records, report_context)


def _rows(records: Sequence[WeatherSnapshot]):
    return [
        [
            item.city,
            item.date_label,
            _number(item.temperature_c),
            item.condition,
            f"{item.humidity_percent}%",
            _number(item.wind_speed_mps),
            "是" if item.rain_expected else "否",
            item.advice or build_travel_advice(
                item.temperature_c,
                item.condition,
                item.humidity_percent,
                item.wind_speed_mps,
                item.rain_expected,
            ),
            item.provider,
        ]
        for item in records
    ]


def build_travel_advice(
    temperature_c: float,
    condition: str,
    humidity_percent: int,
    wind_speed_mps: float,
    rain_expected: bool,
) -> str:
    """根据可验证的天气字段生成简短、完整且逐日可用的出行建议。"""

    suggestions = []
    if rain_expected or "雨" in condition:
        suggestions.append("携带雨伞或轻便雨衣，注意路面湿滑")
    if "雷" in condition:
        suggestions.append("雷电时减少空旷地带和高处停留")
    if temperature_c >= 32:
        suggestions.append("做好防晒并及时补水")
    elif temperature_c <= 10:
        suggestions.append("穿保暖外套")
    elif temperature_c <= 20:
        suggestions.append("带一件薄外套")
    else:
        suggestions.append("适合轻便、透气的分层穿着")
    if wind_speed_mps >= 8:
        suggestions.append("风力较强，优先穿防风外层并留意高空坠物")
    elif wind_speed_mps >= 5:
        suggestions.append("可准备防风外套")
    if humidity_percent >= 80:
        suggestions.append("湿度较高，可准备速干衣物")
    return "；".join(suggestions) + "。"


def _build_markdown(
    records: Sequence[WeatherSnapshot],
    report_context: Optional[WeatherReportContext],
) -> WeatherArtifact:
    headers = ["城市", "日期", "温度（℃）", "天气", "湿度", "风速（m/s）", "有雨", "建议", "数据源"]
    lines = [
        "# 天气报告",
        "",
        *_context_markdown(report_context),
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in _rows(records):
        lines.append("| " + " | ".join(_escape_markdown(str(value)) for value in row) + " |")
    lines.extend(("", "## 出行清单", ""))
    for item, reason in _packing_items(records):
        lines.append(f"- **{item}**：{reason}")
    lines.extend(("", "> 本报告由天气查询 Agent 根据查询时返回的数据生成。", ""))
    return WeatherArtifact(
        _filename(records, "md"),
        "text/markdown",
        "\n".join(lines).encode("utf-8"),
    )


def _build_docx(
    records: Sequence[WeatherSnapshot],
    report_context: Optional[WeatherReportContext],
) -> WeatherArtifact:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    title = document.add_heading("天气报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(26, 68, 109)
    document.add_paragraph("由天气查询 Agent 根据最近一次实时查询结果生成。")
    for label, value in _context_rows(report_context):
        document.add_paragraph(f"{label}：{value}")

    headers = ["城市", "日期", "温度℃", "天气", "湿度", "风速m/s", "有雨", "建议", "数据源"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in _rows(records):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    document.add_heading("出行清单", level=1)
    for item, reason in _packing_items(records):
        document.add_paragraph(f"{item}：{reason}", style="List Bullet")
    output = BytesIO()
    document.save(output)
    return WeatherArtifact(
        _filename(records, "docx"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        output.getvalue(),
    )


def _build_xlsx(
    records: Sequence[WeatherSnapshot],
    report_context: Optional[WeatherReportContext],
) -> WeatherArtifact:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "天气报告"
    headers = ["城市", "日期", "温度（℃）", "天气", "湿度", "风速（m/s）", "是否有雨", "建议", "数据源"]
    sheet.append(headers)
    for row in _rows(records):
        sheet.append([_excel_safe(value) for value in row])

    header_fill = PatternFill("solid", fgColor="1A446D")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    widths = (12, 23, 12, 12, 10, 14, 10, 38, 18)
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    checklist = workbook.create_sheet("出行清单")
    checklist.append(["出行清单", "说明"])
    for label, value in _context_rows(report_context):
        checklist.append([_excel_safe(label), _excel_safe(value)])
    checklist.append([])
    checklist.append(["物品 / 准备", "天气依据"])
    for item, reason in _packing_items(records):
        checklist.append([_excel_safe(item), _excel_safe(reason)])
    for cell in checklist[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
    checklist.freeze_panes = "A2"
    checklist.column_dimensions["A"].width = 24
    checklist.column_dimensions["B"].width = 64
    for row in checklist.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    output = BytesIO()
    workbook.save(output)
    return WeatherArtifact(
        _filename(records, "xlsx"),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        output.getvalue(),
    )


def _build_pdf(
    records: Sequence[WeatherSnapshot],
    report_context: Optional[WeatherReportContext],
) -> WeatherArtifact:
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title="天气报告",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName=PDF_FONT_NAME,
        fontSize=20,
        textColor=colors.HexColor("#1A446D"),
        alignment=TA_CENTER,
    )
    body_style = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName=PDF_FONT_NAME,
        fontSize=8,
        leading=11,
    )
    header_style = ParagraphStyle(
        "ChineseTableHeader",
        parent=body_style,
        textColor=colors.white,
    )
    headers = ["城市", "日期", "温度℃", "天气", "湿度", "风速m/s", "有雨", "建议", "数据源"]
    table_data = [
        [Paragraph(value, header_style) for value in headers],
        *[
            [Paragraph(_escape_xml(str(value)), body_style) for value in row]
            for row in _rows(records)
        ],
    ]
    table = Table(
        table_data,
        repeatRows=1,
        colWidths=[22 * mm, 18 * mm, 20 * mm, 22 * mm, 18 * mm, 22 * mm, 16 * mm, 70 * mm, 35 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A446D")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B9C8D6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F7FA")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story = [Paragraph("天气报告", title_style), Spacer(1, 3 * mm)]
    for label, value in _context_rows(report_context):
        story.append(Paragraph(f"{_escape_xml(label)}：{_escape_xml(value)}", body_style))
    story.extend([Spacer(1, 3 * mm), table, Spacer(1, 4 * mm)])
    story.append(Paragraph("出行清单", title_style))
    for item, reason in _packing_items(records):
        story.append(
            Paragraph(
                f"• {_escape_xml(item)}：{_escape_xml(reason)}",
                body_style,
            )
        )
    story.extend(
        [
            Spacer(1, 3 * mm),
            Paragraph("本报告由天气查询 Agent 根据查询时返回的数据生成。", body_style),
        ]
    )
    document.build(story)
    return WeatherArtifact(_filename(records, "pdf"), "application/pdf", output.getvalue())


def _filename(records: Sequence[WeatherSnapshot], extension: str) -> str:
    unique_cities = tuple(dict.fromkeys(item.city for item in records))
    city_part = "-".join(
        re.sub(r"[^\w\u4e00-\u9fff-]", "", city) for city in unique_cities
    )
    safe_city_part = city_part[:48] or "weather"
    return f"天气报告-{safe_city_part}.{extension}"


def _number(value: float) -> str:
    return f"{value:.1f}".rstrip("0").rstrip(".")


def _escape_markdown(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")


def _excel_safe(value):
    if isinstance(value, str) and value.startswith(("=", "+", "-", "@")):
        return f"'{value}"
    return value


def _escape_xml(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _context_rows(
    report_context: Optional[WeatherReportContext],
) -> Tuple[Tuple[str, str], ...]:
    if report_context is None:
        return ()
    rows = []
    if report_context.origin:
        rows.append(("出发地", report_context.origin))
    if report_context.total_days:
        rows.append(("总行程天数", f"{report_context.total_days} 天"))
        rows.append(
            (
                "日期分配方式",
                "按目的地顺序逐日安排"
                if report_context.ordered
                else "每个目的地均列出全部行程日",
            )
        )
    return tuple(rows)


def _context_markdown(
    report_context: Optional[WeatherReportContext],
) -> Tuple[str, ...]:
    rows = _context_rows(report_context)
    if not rows:
        return ()
    return tuple([*(f"- **{label}**：{value}" for label, value in rows), ""])


def _packing_items(
    records: Sequence[WeatherSnapshot],
) -> Tuple[Tuple[str, str], ...]:
    items = []
    if any(item.rain_expected or "雨" in item.condition for item in records):
        items.append(("雨伞或轻便雨衣", "行程中存在降雨天气，并注意准备防滑鞋履"))
    if any("雷" in item.condition for item in records):
        items.append(("雷雨备选安排", "雷电时减少空旷地带、高处和水边活动"))
    if any(item.temperature_c >= 30 for item in records):
        items.append(("防晒用品与饮用水", "部分行程日气温较高"))
    if any(item.temperature_c <= 20 for item in records):
        items.append(("薄外套或保暖外层", "部分行程日气温偏低"))
    if any(item.wind_speed_mps >= 5 for item in records):
        items.append(("防风外套", "部分行程日风力较明显"))
    if any(item.humidity_percent >= 80 for item in records):
        items.append(("速干衣物", "部分行程日湿度较高"))
    if not items:
        items.append(("轻便分层衣物", "当前预报未显示明显雨、强风或极端温度"))
    return tuple(items)
