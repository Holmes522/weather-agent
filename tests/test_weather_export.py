from io import BytesIO

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from weather_export import (
    WeatherReportContext,
    WeatherSnapshot,
    build_weather_document,
    parse_export_request,
    weather_query_text,
)


SAMPLE = WeatherSnapshot(
    city="深圳",
    date_label="明天",
    provider="Open-Meteo",
    temperature_c=25.0,
    condition="雨",
    humidity_percent=80,
    wind_speed_mps=2.0,
    rain_expected=True,
    advice="明天可能有雨，建议携带雨具。",
)


def test_export_intent_recognizes_common_formats_and_excel_typo():
    assert parse_export_request("把刚才天气保存成 Word").format == "docx"
    assert parse_export_request("导出为 execl 表格").format == "xlsx"
    assert parse_export_request("生成PDF给我").format == "pdf"
    assert parse_export_request("整理成 md").format == "md"
    assert parse_export_request("深圳明天天气怎么样").requested is False


def test_export_directive_removal_keeps_itinerary_details_after_the_format():
    cleaned = weather_query_text(
        "依次查看北京和深圳天气，结果输出Excel表，我住在杭州，每个城市待一天"
    )

    assert "北京和深圳天气" in cleaned
    assert "我住在杭州" in cleaned
    assert "每个城市待一天" in cleaned
    assert "输出Excel" not in cleaned


def test_markdown_export_contains_structured_weather_fields():
    artifact = build_weather_document((SAMPLE,), "md")

    text = artifact.content.decode("utf-8")
    assert artifact.filename.endswith(".md")
    assert "深圳" in text
    assert "80%" in text
    assert "Open-Meteo" in text


def test_docx_export_can_be_reopened():
    artifact = build_weather_document((SAMPLE,), "docx")

    document = Document(BytesIO(artifact.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = " ".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert artifact.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "天气报告" in text
    assert "出行清单" in text
    assert "雨伞或轻便雨衣" in text
    assert "深圳" in table_text
    assert "建议携带雨具" in table_text


def test_xlsx_export_can_be_reopened_without_formulas():
    artifact = build_weather_document(
        (SAMPLE,),
        "xlsx",
        WeatherReportContext(origin="杭州", total_days=1, ordered=True),
    )

    workbook = load_workbook(BytesIO(artifact.content), data_only=False)
    sheet = workbook.active
    values = [cell.value for row in sheet.iter_rows() for cell in row]
    assert sheet.title == "天气报告"
    assert "深圳" in values
    assert sheet["H2"].value
    assert "出行清单" in workbook.sheetnames
    checklist_values = [
        cell.value
        for row in workbook["出行清单"].iter_rows()
        for cell in row
    ]
    assert "杭州" in checklist_values
    assert any("雨" in str(value) for value in checklist_values)
    assert not any(isinstance(value, str) and value.startswith("=") for value in values)


def test_export_generates_advice_when_snapshot_has_none():
    no_advice = WeatherSnapshot(
        city="北京",
        date_label="第3天（2026-09-01）",
        provider="Open-Meteo",
        temperature_c=34.0,
        condition="晴",
        humidity_percent=45,
        wind_speed_mps=2.0,
        rain_expected=False,
    )

    artifact = build_weather_document((no_advice,), "md")
    text = artifact.content.decode("utf-8")

    assert "防晒" in text
    assert "出行清单" in text


def test_xlsx_export_escapes_formula_like_external_text():
    unsafe = WeatherSnapshot(
        city="=2+2",
        date_label="明天",
        provider="Open-Meteo",
        temperature_c=25.0,
        condition="晴",
        humidity_percent=50,
        wind_speed_mps=2.0,
        rain_expected=False,
    )

    artifact = build_weather_document((unsafe,), "xlsx")
    workbook = load_workbook(BytesIO(artifact.content), data_only=False)
    cell = workbook.active["A2"]

    assert cell.data_type != "f"
    assert cell.value == "'=2+2"


def test_pdf_export_can_be_reopened_and_contains_one_page():
    artifact = build_weather_document((SAMPLE,), "pdf")

    reader = PdfReader(BytesIO(artifact.content))
    assert artifact.mimetype == "application/pdf"
    assert len(reader.pages) == 1
